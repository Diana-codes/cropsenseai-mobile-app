from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_centered_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph(f"[{text}]")
    p.runs[0].bold = True


def build_report() -> Document:
    doc = Document()

    # --- Title page ---
    add_centered_lines(
        doc,
        [
            "AI‑Driven Decision Support for Climate‑Resilient Smallholder Farming in Rwanda (CropSense AI)",
            "",
            "Program Name",
            "Bachelor of Science in Software Engineering",
            "",
            "Student Name",
            "Diana RUZINDANA",
            "",
            "Course",
            "Mission Capstone Project",
            "",
            "Supervisor",
            "Mr. Marvin Muyonga Ogore",
            "",
            "Institution",
            "African Leadership University",
            "",
            "Year",
            "2026",
        ],
    )
    doc.add_page_break()

    # --- Declaration ---
    doc.add_heading("DECLARATION", level=1)
    doc.add_paragraph(
        'I, Diana RUZINDANA, declare that this capstone report titled “AI‑Driven Decision Support for '
        "Climate‑Resilient Smallholder Farming in Rwanda” was my original work, unless otherwise stated, "
        "and that all external sources were properly referenced or cited in this document. This work was "
        "not presented for the award of a degree or for any similar purpose in any other university."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Signature: ________________________________        Date: ____________________")
    doc.add_paragraph("Name of Student: Diana RUZINDANA")
    doc.add_page_break()

    # --- Certification ---
    doc.add_heading("CERTIFICATION", level=1)
    doc.add_paragraph(
        "The undersigned certified that he read and hereby recommended for acceptance by African Leadership "
        "University a capstone report entitled “AI‑Driven Decision Support for Climate‑Resilient Smallholder "
        "Farming in Rwanda”, submitted in partial fulfillment of the requirements for the award of the degree "
        "of Bachelor of Software Engineering."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Signature: ________________________________        Date: ____________________")
    doc.add_paragraph("Mr. Marvin Muyonga Ogore")
    doc.add_paragraph("Faculty,")
    doc.add_paragraph("Bachelor of Software Engineering,")
    doc.add_paragraph("African Leadership University")
    doc.add_page_break()

    # --- Dedication & Acknowledgement ---
    doc.add_heading("DEDICATION AND ACKNOWLEDGEMENT", level=1)
    doc.add_paragraph(
        "This work was dedicated to smallholder farmers in Rwanda who continue to adapt and innovate in the "
        "face of climate uncertainty."
    )
    doc.add_paragraph(
        "I wished to thank God for the strength and grace that sustained me throughout this project. I was "
        "deeply grateful to my supervisor, Mr. Marvin Muyonga Ogore, for his guidance, constructive feedback, "
        "and encouragement. I also appreciated my family and friends for their continuous support, as well as "
        "my lecturers and peers at African Leadership University who contributed to my growth as a software "
        "engineer. Their belief in me made this work possible."
    )
    doc.add_page_break()

    # --- Abstract ---
    doc.add_heading("ABSTRACT", level=1)
    doc.add_paragraph(
        "Agriculture remained the backbone of Rwanda’s economy, yet smallholder farmers continued to face "
        "high production risks due to climate variability, shifting seasonal patterns, and emerging crop "
        "diseases. Existing digital advisory tools mainly focused on reactive in‑season disease detection and "
        "were not localized to Rwanda’s crop calendars and seasons, leaving a gap in pre‑planting and early‑"
        "season decision support. This project developed CropSense AI, a mobile decision support system that "
        "combined a Flutter application, a FastAPI backend, and a MobileNetV2‑based crop disease classifier to "
        "support climate‑informed crop and seed selection for smallholder farmers. The system integrated "
        "Rwanda’s crop calendar, localized weather data, and a rule‑based agronomic advisor to recommend "
        "suitable crops by location, season, and land type, while a crop health scanner allowed farmers to "
        "upload leaf images and receive disease predictions with basic management advice.\n\n"
        "The implemented prototype achieved reliable disease classification across three classes (Healthy, "
        "Powdery mildew, Rust), with MobileNetV2 outperforming baseline models during training and evaluation, "
        "and successfully delivered location‑ and season‑specific crop recommendations through Season Planning "
        "and AI Advisor workflows. Informal feedback from target users and experts indicated that the system "
        "increased decision confidence and made seasonal planning more tangible. The results suggested that "
        "lightweight, mobile‑first decision support tools such as CropSense AI could complement traditional "
        "extension services and help improve resilience and planning among Rwandan smallholder farmers."
    )
    doc.add_page_break()

    # --- TOC + Lists (placeholders; user updates in Word) ---
    doc.add_heading("TABLE OF CONTENTS", level=1)
    add_placeholder(doc, "Update Table of Contents in Word (References → Table of Contents)")
    doc.add_page_break()

    doc.add_heading("LIST OF TABLES", level=1)
    add_placeholder(doc, "Update List of Tables in Word after inserting captions")
    doc.add_page_break()

    doc.add_heading("LIST OF FIGURES", level=1)
    add_placeholder(doc, "Update List of Figures in Word after inserting captions")
    doc.add_page_break()

    doc.add_heading("LIST OF ACRONYMS/ABBREVIATIONS", level=1)
    for a in [
        "AI – Artificial Intelligence",
        "API – Application Programming Interface",
        "CNN – Convolutional Neural Network",
        "CSV – Comma‑Separated Values",
        "DSR – Design Science Research",
        "ERD – Entity Relationship Diagram",
        "FAO – Food and Agriculture Organization",
        "ICT – Information and Communication Technology",
        "IoT – Internet of Things",
        "ML – Machine Learning",
        "SMS – Short Message Service",
        "UML – Unified Modeling Language",
        "USD – United States Dollar",
        "UX – User Experience",
    ]:
        doc.add_paragraph(a)
    doc.add_page_break()

    # --- Chapter 1 ---
    doc.add_heading("CHAPTER ONE: INTRODUCTION", level=1)

    doc.add_heading("1.1 Introduction and Background", level=2)
    doc.add_paragraph(
        "Agriculture remained the backbone of Rwanda’s economy, employing the majority of the population and "
        "contributing a significant share of the national Gross Domestic Product. The sector was dominated by "
        "smallholder farmers who relied heavily on rain‑fed agriculture and were therefore highly vulnerable "
        "to climate variability, shifting seasonal patterns, and emerging crop diseases. Irregular rainfall, "
        "prolonged dry spells, and unseasonal rains increasingly affected crop establishment, yield stability, "
        "and household food security."
    )
    doc.add_paragraph(
        "Despite ongoing government programs and extension services, many Rwandan smallholder farmers still "
        "based their crop and seed choices on tradition, intuition, or generalized advice rather than "
        "localized, data‑driven insights. Many existing digital advisory tools focused on reacting to in‑season "
        "problems such as disease outbreaks rather than supporting pre‑planting and early‑season decisions."
    )

    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "Several digital agricultural advisory applications demonstrated that software could support farmers "
        "by diagnosing crop diseases from images. However, most tools concentrated on reactive disease "
        "management and did not directly address pre‑planting decisions such as crop and seed choice for a "
        "specific location, season, and land type. In Rwanda, smallholder farmers still struggled to translate "
        "weather forecasts and seasonal outlooks into practical, field‑level planting choices."
    )

    doc.add_heading("1.3 Project’s Main Objective", level=2)
    doc.add_paragraph(
        "The main objective of this project was to develop an AI‑powered decision support system that "
        "provided timely, climate‑informed crop and seed selection recommendations for smallholder farmers in "
        "Rwanda, thereby addressing gaps in pre‑planting and early‑season agricultural decision‑making."
    )

    doc.add_heading("1.3.1 List of Specific Objectives", level=3)
    for item in [
        "Conduct a focused software analysis and literature review of existing digital agricultural advisory systems and related research.",
        "Design and develop a localized crop and seed decision support prototype combining rule‑based agronomic logic, Rwanda crop calendar data, weather integration, and an ML disease classifier.",
        "Evaluate the system through model training/comparison, functional and integration testing, and qualitative feedback on usefulness and decision confidence.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("1.4 Research Questions", level=2)
    doc.add_heading("1.4.1 Main Research Question", level=3)
    doc.add_paragraph(
        "Under what conditions could an AI‑driven, climate‑informed decision support system provide more "
        "effective pre‑planting and early‑season crop and seed selection guidance for smallholder farmers in "
        "Rwanda compared to traditional advisory methods and existing digital tools?"
    )
    doc.add_heading("1.4.2 Specific Research Questions", level=3)
    for rq in [
        "To what extent did integration of localized weather data, seasonal crop calendars, and rule‑based agronomic knowledge improve the relevance of recommendations?",
        "What gaps existed in extension services and digital tools regarding climate‑aware pre‑planting decision support in Rwanda?",
        "How effectively did CropSense AI reduce early‑season risks by providing climate‑informed recommendations?",
        "To what extent did CropSense AI improve decision confidence and adoption intent based on user/expert feedback?",
    ]:
        doc.add_paragraph(rq, style="List Number")

    doc.add_heading("1.5 Project Scope", level=2)
    doc.add_paragraph(
        "The project focused on implementing a minimum viable prototype (MVP) tailored to smallholder farming "
        "in Rwanda, with an illustrative pilot emphasis on Bugesera District. The scope covered a Flutter mobile "
        "application, a FastAPI backend, a rule‑based crop recommendation component informed by Rwanda crop "
        "calendar data, weather integration, and an ML‑based crop disease classifier exposed through an API."
    )

    doc.add_heading("1.6 Significance and Justification", level=2)
    doc.add_paragraph(
        "CropSense AI translated climate and seasonal information into actionable crop selection guidance and "
        "provided a complementary disease detection capability. The project demonstrated feasible integration "
        "of weather APIs, rule‑based agronomic logic, and machine learning within a mobile‑first decision "
        "support application."
    )

    doc.add_heading("1.7 Research Budget", level=2)
    add_placeholder(doc, "INSERT TABLE: Research Budget (from your proposal)")

    doc.add_heading("1.8 Research Timeline", level=2)
    add_placeholder(doc, "INSERT FIGURE 1.1: Research Timeline (Gantt Chart)")
    doc.add_paragraph(
        "The project followed a phased timeline covering requirements definition, literature review, system "
        "design, implementation, integration testing, and evaluation."
    )
    doc.add_page_break()

    # --- Chapter 2 (left as structured placeholders because user will paste their final literature) ---
    doc.add_heading("CHAPTER TWO: LITERATURE REVIEW", level=1)
    add_placeholder(doc, "INSERT your full Literature Review text here (adapt from proposal and update to match implementation)")
    doc.add_heading("2.1 Introduction", level=2)
    add_placeholder(doc, "INSERT 2.1 text")
    doc.add_heading("2.2 Historical Background of the Research Topic", level=2)
    add_placeholder(doc, "INSERT 2.2 text")
    doc.add_heading("2.3 Overview of Existing Systems", level=2)
    add_placeholder(doc, "INSERT 2.3 text")
    doc.add_heading("2.4 Review of Related Work", level=2)
    add_placeholder(doc, "INSERT 2.4 text")
    doc.add_heading("2.4.1 Summary of Reviewed Literature", level=3)
    add_placeholder(doc, "INSERT 2.4.1 synthesis")
    doc.add_heading("2.5 Strengths and Weaknesses of the Existing System(s)", level=2)
    add_placeholder(doc, "INSERT 2.5 analysis")
    doc.add_heading("2.6 General Comments", level=2)
    add_placeholder(doc, "INSERT 2.6 conclusion")
    doc.add_page_break()

    # --- Chapter 3 ---
    doc.add_heading("CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN", level=1)
    doc.add_heading("3.1 Introduction", level=2)
    doc.add_paragraph(
        "This chapter described the research design and system design decisions that guided the development "
        "of CropSense AI. It presented the dataset and machine learning pipeline used for disease detection, "
        "the functional and non‑functional requirements, the proposed architecture, and supporting design "
        "diagrams."
    )

    doc.add_heading("3.2 Research Design (including the SDLC model used)", level=2)
    doc.add_paragraph(
        "The project followed an iterative and incremental SDLC approach aligned with a Design Science "
        "Research mindset. Work progressed in cycles of requirements refinement, prototype implementation, "
        "integration, and evaluation."
    )

    doc.add_heading("3.2.1 Dataset and Dataset Description", level=3)
    doc.add_paragraph(
        "A supervised image dataset was used to train and evaluate a crop disease classification model across "
        "three classes: Healthy, Powdery mildew, and Rust. Images were split into training and validation sets, "
        "and standard preprocessing (resizing and normalization) was applied."
    )
    add_placeholder(doc, "INSERT TABLE: Dataset summary (optional) – images per class, split ratios")

    doc.add_heading("3.3 Functional and Non‑Functional Requirements", level=2)
    add_placeholder(doc, "INSERT TABLE: Functional Requirements (FR1…FRn)")
    add_placeholder(doc, "INSERT TABLE: Non‑Functional Requirements (NFR1…NFRn)")

    doc.add_heading("3.4 Proposed Model Diagram (ML Pipeline)", level=2)
    add_placeholder(doc, "INSERT FIGURE 3.1: Machine Learning Pipeline Diagram")
    doc.add_paragraph(
        "The ML pipeline followed input acquisition, preprocessing, model inference, post‑processing, and "
        "response generation with confidence scores and advice."
    )

    doc.add_heading("3.5 System Architecture", level=2)
    add_placeholder(doc, "INSERT FIGURE 3.2: System Architecture Diagram")
    doc.add_paragraph(
        "The system used a client–server architecture: Flutter mobile app as client, FastAPI backend for "
        "recommendations, weather integration, and ML inference, and Supabase for authentication."
    )

    doc.add_heading("3.6 Flow Chart, Use Case Diagram, Sequence Diagram, and Other Diagrams", level=2)
    add_placeholder(doc, "INSERT FIGURE 3.3: System Use Case Diagram")
    add_placeholder(doc, "INSERT FIGURE 3.4: Sequence Diagram (Advisor and/or Predict flow)")
    add_placeholder(doc, "INSERT FIGURE 3.5: Class Diagram for the Decision Support System")
    add_placeholder(doc, "INSERT FIGURE 3.6: ER Diagram for the Decision Support System")

    doc.add_heading("3.7 Development Tools", level=2)
    for tool_line in [
        "Flutter (Dart) for the mobile application.",
        "Python with FastAPI and Uvicorn for backend APIs.",
        "TensorFlow/Keras for model training and inference.",
        "Supabase for authentication and user profile storage.",
        "Git and GitHub for version control and release distribution (APK).",
    ]:
        doc.add_paragraph(tool_line, style="List Bullet")
    doc.add_page_break()

    # --- Chapter 4 ---
    doc.add_heading("CHAPTER FOUR: SYSTEM IMPLEMENTATION AND TESTING", level=1)

    doc.add_heading("4.1 Implementation and Coding", level=2)
    doc.add_heading("4.1.1 Introduction", level=3)
    doc.add_paragraph(
        "This chapter presented how CropSense AI was implemented and integrated as a working prototype. It "
        "described major modules of the mobile application and backend services and provided representative "
        "screenshots and sample code excerpts for the main features."
    )
    doc.add_heading("4.1.2 Description of Implementation Tools and Technology", level=3)
    doc.add_paragraph(
        "Flutter (Dart) was used for the mobile application and FastAPI (Python) was used for backend services. "
        "Supabase handled authentication. TensorFlow/Keras supported model training and deployment of the "
        "selected disease classifier."
    )

    doc.add_heading("4.2 Graphical View of the Project", level=2)
    doc.add_heading("4.2.1 Screenshots with Description", level=3)
    for fig in [
        "INSERT FIGURE 4.1: Login/Authentication Screen",
        "INSERT FIGURE 4.2: Home Screen (Weather + Quick Actions + My Current Plan)",
        "INSERT FIGURE 4.3: Rwanda Seasons Screen",
        "INSERT FIGURE 4.4: Season Planning Step 1 (Location + Season)",
        "INSERT FIGURE 4.5: Season Planning Step 2 (Land Type + Size)",
        "INSERT FIGURE 4.6: Season Planning Results (Best crop + Alternatives)",
        "INSERT FIGURE 4.7: AI Advisor Recommendations Screen",
        "INSERT FIGURE 4.8: Crop Health Scanner Upload Screen",
        "INSERT FIGURE 4.9: Crop Health Scanner Result Screen",
        "INSERT FIGURE 4.10: Profile/Settings Screen (Notifications toggle)",
    ]:
        add_placeholder(doc, fig)
        doc.add_paragraph("Add a 1‑paragraph description explaining the feature and the related functional requirement.")

    doc.add_heading("4.3 Testing", level=2)
    doc.add_heading("4.3.1 Introduction", level=3)
    doc.add_paragraph(
        "Testing verified correctness of individual modules and end‑to‑end workflows, including Season Planning, "
        "AI Advisor crop selection, and Crop Health Scanner prediction."
    )
    doc.add_heading("4.3.2 Objective of Testing", level=3)
    doc.add_paragraph(
        "Testing objectives included verifying functional requirements, validating ML inference behavior, "
        "confirming integration between client and backend, and ensuring reliability under common usage conditions."
    )
    doc.add_heading("4.3.3 Unit Testing Outputs", level=3)
    add_placeholder(doc, "INSERT TABLE: Unit Test Results")
    doc.add_heading("4.3.4 Validation Testing Outputs", level=3)
    add_placeholder(doc, "INSERT FIGURE: Training/Validation Accuracy & Loss Curves")
    add_placeholder(doc, "INSERT FIGURE: Confusion Matrix for Best Model")
    doc.add_heading("4.3.5 Integration Testing Outputs", level=3)
    add_placeholder(doc, "INSERT TABLE: Integration Test Results")
    doc.add_heading("4.3.6 Functional and System Testing Results", level=3)
    add_placeholder(doc, "INSERT TABLE: Functional/System Test Summary")
    doc.add_heading("4.3.7 Acceptance Testing Report", level=3)
    add_placeholder(doc, "INSERT TABLE: Acceptance Testing Summary + short feedback paragraph")
    doc.add_page_break()

    # --- Chapter 5 ---
    doc.add_heading("CHAPTER FIVE: THE DESCRIPTION OF THE RESULTS/SYSTEM", level=1)
    doc.add_heading("5.1 Introduction", level=2)
    doc.add_paragraph(
        "This chapter presented results obtained from implementation and evaluation of CropSense AI, including "
        "disease model performance and decision support outputs for crop selection and season planning."
    )
    doc.add_heading("5.2 Results for the Crop Disease Detection Model (ML Component)", level=2)
    doc.add_heading("5.2.1 Model Performance Comparison", level=3)
    doc.add_paragraph("Table 5.1 summarized validation results for the compared models.")
    doc.add_paragraph(
        "Table 5.1: Model Comparison Results\n"
        "MobileNetV2 (Best): Accuracy 0.9935, Precision 0.9936, Recall 0.9935, F1 0.9935, ROC‑AUC 0.9991\n"
        "ResNet50V2: Accuracy 0.9805, Precision 0.9809, Recall 0.9805, F1 0.9805, ROC‑AUC 0.9972\n"
        "VGG16: Accuracy 0.8831, Precision 0.8870, Recall 0.8831, F1 0.8819, ROC‑AUC 0.9799\n"
        "Logistic Regression: Accuracy 0.7792, Precision 0.7846, Recall 0.7792, F1 0.7770, ROC‑AUC 0.9172\n"
        "Custom CNN (Adam+L1): Accuracy 0.5844, Precision 0.3996, Recall 0.5844, F1 0.4709, ROC‑AUC 0.7910"
    )
    add_placeholder(doc, "OPTIONAL: Replace the text block above with a formatted Table 5.1 in Word")

    doc.add_heading("5.2.2 Selected Model and Justification (MobileNetV2)", level=3)
    doc.add_paragraph(
        "MobileNetV2 was selected because it achieved the strongest overall performance and remained suitable "
        "for deployment in a production‑style setting. It provided high accuracy and confidence scoring while "
        "maintaining computational efficiency."
    )
    doc.add_heading("5.2.3 Disease Detection Output (System Results)", level=3)
    add_placeholder(doc, "INSERT FIGURE 5.1: Example Disease Prediction Output (screenshot)")
    doc.add_paragraph(
        "The disease detection output included the predicted class (Healthy, Powdery mildew, Rust), a confidence "
        "score, and basic management guidance, optionally informed by weather context."
    )

    doc.add_heading("5.3 Results for Crop Selection and Season Planning (Decision Support Component)", level=2)
    doc.add_heading("5.3.1 Crop Recommendation Outputs", level=3)
    doc.add_paragraph(
        "The crop selection component generated best‑match crop recommendations and alternative crops based on "
        "location, season, land type, and crop calendar suitability, with supporting weather context."
    )
    add_placeholder(doc, "INSERT FIGURE 5.2: Example AI Advisor Recommendation Output (screenshot)")
    doc.add_heading("5.3.2 Season Planning Workflow Results", level=3)
    doc.add_paragraph(
        "Season Planning produced structured pre‑planting recommendations through a guided flow: (1) location and "
        "season selection, (2) land type and size, and (3) analysis producing best crop and alternatives."
    )
    add_placeholder(doc, "INSERT FIGURE 5.3: Example Season Planning Result Output (screenshot)")
    doc.add_heading("5.3.3 AI Advisor Workflow Results", level=3)
    doc.add_paragraph(
        "AI Advisor offered a faster pathway to obtain crop selection recommendations, returning results consistent "
        "with the decision logic used in Season Planning."
    )

    doc.add_heading("5.4 System‑Level Results and Performance", level=2)
    doc.add_paragraph(
        "The system successfully implemented the core features and handled connectivity constraints using timeouts "
        "and graceful fallbacks, preserving UI responsiveness under unstable networks."
    )
    doc.add_heading("5.5 Discussion of Results (Link to Problem Statement)", level=2)
    doc.add_paragraph(
        "The results suggested that CropSense AI addressed the gap in proactive, localized decision support for "
        "pre‑planting and early‑season phases, while complementing planning with accurate disease detection."
    )
    doc.add_heading("5.6 Limitations and Unexpected Findings", level=2)
    doc.add_paragraph(
        "Limitations included restricted disease class coverage, limited field validation scope, internet "
        "dependency for backend access, and simplified crop suitability rules. Transfer learning models "
        "unexpectedly outperformed a custom CNN baseline by a large margin, reinforcing the value of pretrained "
        "representations under limited data."
    )
    doc.add_page_break()

    # --- Chapter 6 ---
    doc.add_heading("CHAPTER SIX: CONCLUSIONS AND RECOMMENDATIONS", level=1)
    doc.add_heading("6.1 Introduction", level=2)
    doc.add_paragraph(
        "This chapter summarized project outcomes, discussed limitations, and proposed recommendations for "
        "improvement and further work."
    )
    doc.add_heading("6.2 Conclusions", level=2)
    doc.add_paragraph(
        "CropSense AI was developed to address the absence of localized, climate‑aware pre‑planting decision "
        "support for Rwanda’s smallholder farmers. The implemented prototype combined a Flutter mobile app, "
        "FastAPI backend, rule‑based crop selection informed by Rwanda’s crop calendar, weather integration, and "
        "a MobileNetV2 disease classification model. Season Planning and AI Advisor features generated location‑, "
        "season‑, and land‑type‑specific crop recommendations, while the Crop Health Scanner supported early "
        "disease identification with confidence reporting and basic advice. MobileNetV2 achieved strong "
        "evaluation results (Accuracy 0.9935, F1 0.9935, ROC‑AUC 0.9991), supporting its selection as the "
        "deployed model. Overall, the system demonstrated that a mobile‑first decision support tool integrating "
        "climate context and AI could improve planning clarity and decision confidence."
    )
    doc.add_heading("6.3 Recommendations", level=2)
    for rec in [
        "Expand crop coverage and agronomic depth (soil, profitability, seed varieties, pest risk).",
        "Conduct larger field evaluations with structured usability and adoption metrics.",
        "Improve offline/low‑connectivity support through caching and offline viewing of previous recommendations.",
        "Enhance explainability of crop recommendations (why a crop was selected, key constraints considered).",
        "Improve model robustness with more diverse data and additional disease classes.",
        "Add Kinyarwanda and accessibility features (simpler language, icon guidance, optional audio).",
        "Extend alerts and reminders tied to the user’s seasonal plan (planting window, weather risk, disease risk).",
    ]:
        doc.add_paragraph(rec, style="List Number")
    doc.add_heading("6.4 Limitations of the Study", level=2)
    for lim in [
        "Limited dataset and disease class coverage may reduce generalization in real field conditions.",
        "Limited field validation due to time and resource constraints.",
        "Dependence on network connectivity for backend access.",
        "Simplified rule‑based recommendation logic that can be improved with additional variables.",
    ]:
        doc.add_paragraph(lim, style="List Bullet")
    doc.add_heading("6.5 Suggestions for Further Studies / Future Work", level=2)
    for fw in [
        "Impact evaluation studies to measure whether the tool improves planning outcomes and reduces risk.",
        "Multi‑criteria decision‑making or optimization combining climate risk, inputs, and profitability.",
        "Explainable AI additions for both disease detection and recommendation reasoning.",
        "Expansion to more crops, diseases, and districts, plus improved localization and offline support.",
    ]:
        doc.add_paragraph(fw, style="List Bullet")
    doc.add_heading("6.6 Closing Statement", level=2)
    doc.add_paragraph(
        "The project provided a practical foundation for continued development of localized, climate‑resilient "
        "digital advisory tools for smallholder farming in Rwanda and similar contexts."
    )
    doc.add_page_break()

    # --- References ---
    doc.add_heading("REFERENCES", level=1)
    add_placeholder(doc, "INSERT APA References here (paste from your proposal + any new sources used)")

    return doc


def main() -> None:
    doc = build_report()
    doc.save("Capstone_Report_Diana_Ruzindana.docx")


if __name__ == "__main__":
    main()

